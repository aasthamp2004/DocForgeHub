import os
import streamlit as st
import requests
import json
from docx import Document
from docx.shared import Pt
from io import BytesIO
import pandas as pd

# Reads BACKEND_URL from env when running in Docker,
# falls back to localhost for local development
BACKEND = os.getenv("BACKEND_URL", "http://127.0.0.1:8000")

st.set_page_config(page_title="DocForge AI", layout="wide", initial_sidebar_state="expanded")

# ─────────────────────────────────────────────────────────────────────────────
# CSS — Navy + Gold theme
# ─────────────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
/* ── Google Font ── */
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');

/* ── Global ── */
html, body, [class*="css"] {
    font-family: 'Inter', sans-serif;
}

/* ── App background ── */
.stApp {
    background: linear-gradient(135deg, #0a0e1a 0%, #0d1425 50%, #0a0e1a 100%);
}

/* ── Sidebar ── */
[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #0d1631 0%, #0a1128 100%) !important;
    border-right: 1px solid #1e2d4a !important;
}
[data-testid="stSidebar"] .stRadio label {
    color: #a0b4cc !important;
    font-size: 0.92rem;
    padding: 6px 0;
}
[data-testid="stSidebar"] .stRadio label:hover {
    color: #f5c842 !important;
}

/* ── Page title ── */
.page-header {
    background: linear-gradient(90deg, #0d1631, #112044);
    border: 1px solid #1e3a6e;
    border-left: 4px solid #f5c842;
    border-radius: 12px;
    padding: 24px 32px;
    margin-bottom: 28px;
}
.page-header h1 {
    font-size: 1.9rem;
    font-weight: 700;
    color: #ffffff;
    margin: 0 0 4px 0;
}
.page-header p {
    color: #7a94b8;
    font-size: 0.88rem;
    margin: 0;
}

/* ── Step cards ── */
.step-card {
    background: linear-gradient(135deg, #0f1a2e 0%, #111f38 100%);
    border: 1px solid #1e3055;
    border-radius: 14px;
    padding: 24px 28px;
    margin-bottom: 20px;
    position: relative;
}
.step-label {
    font-size: 0.7rem;
    font-weight: 700;
    letter-spacing: 2px;
    color: #f5c842;
    text-transform: uppercase;
    margin-bottom: 4px;
}

/* ── Document content card ── */
.doc-card {
    background: linear-gradient(135deg, #0f1a2e 0%, #0d1731 100%);
    border: 1px solid #1e3055;
    border-radius: 14px;
    padding: 32px 36px;
    margin-bottom: 24px;
    box-shadow: 0 4px 24px rgba(0,0,0,0.3);
}
.doc-title {
    font-size: 1.7rem;
    font-weight: 700;
    color: #ffffff;
    border-bottom: 2px solid #f5c842;
    padding-bottom: 12px;
    margin-bottom: 24px;
}
.doc-section-heading {
    font-size: 1.05rem;
    font-weight: 600;
    color: #f5c842;
    margin-top: 24px;
    margin-bottom: 8px;
    padding-left: 12px;
    border-left: 3px solid #f5c842;
}
.doc-paragraph {
    font-size: 0.93rem;
    color: #c8d8ec;
    line-height: 1.85;
    margin-bottom: 10px;
    padding-left: 15px;
}
.bullet-item {
    color: #c8d8ec;
    font-size: 0.93rem;
    line-height: 1.85;
    padding-left: 22px;
    position: relative;
    margin-bottom: 4px;
}
.bullet-item::before {
    content: "▸";
    color: #f5c842;
    position: absolute;
    left: 6px;
    font-size: 0.8rem;
}

/* ── Format badges ── */
.format-badge-excel {
    background: rgba(245,200,66,0.12);
    color: #f5c842;
    border: 1px solid #f5c842;
    border-radius: 20px;
    padding: 3px 12px;
    font-size: 11px;
    font-weight: 700;
    letter-spacing: 0.5px;
}
.format-badge-word {
    background: rgba(59,130,246,0.12);
    color: #60a5fa;
    border: 1px solid #3b82f6;
    border-radius: 20px;
    padding: 3px 12px;
    font-size: 11px;
    font-weight: 700;
    letter-spacing: 0.5px;
}

/* ── History cards ── */
.hist-card {
    background: linear-gradient(135deg, #0f1a2e 0%, #111f38 100%);
    border: 1px solid #1e3055;
    border-radius: 12px;
    padding: 18px 22px;
    margin-bottom: 14px;
    transition: border-color 0.2s;
}
.hist-card:hover { border-color: #f5c842; }
.hist-title {
    font-size: 1rem;
    font-weight: 700;
    color: #e8f0ff;
}
.hist-meta {
    font-size: 0.8rem;
    color: #6b82a8;
    margin-top: 5px;
}

/* ── Streamlit buttons ── */
.stButton > button {
    background: linear-gradient(135deg, #1a3a6e, #1e4494) !important;
    color: #ffffff !important;
    border: 1px solid #2a52a8 !important;
    border-radius: 8px !important;
    font-weight: 600 !important;
    font-size: 0.88rem !important;
    padding: 0.45rem 1.2rem !important;
    transition: all 0.2s !important;
}
.stButton > button:hover {
    background: linear-gradient(135deg, #f5c842, #e6b830) !important;
    color: #0a0e1a !important;
    border-color: #f5c842 !important;
    transform: translateY(-1px);
    box-shadow: 0 4px 12px rgba(245,200,66,0.3) !important;
}

/* ── Primary / download button ── */
.stDownloadButton > button {
    background: linear-gradient(135deg, #f5c842, #e6b830) !important;
    color: #0a0e1a !important;
    border: none !important;
    border-radius: 8px !important;
    font-weight: 700 !important;
    font-size: 0.9rem !important;
}
.stDownloadButton > button:hover {
    background: linear-gradient(135deg, #ffe066, #f5c842) !important;
    box-shadow: 0 4px 16px rgba(245,200,66,0.4) !important;
}

/* ── Text inputs ── */
.stTextInput > div > div > input,
.stTextArea > div > div > textarea {
    background: #0d1631 !important;
    border: 1px solid #1e3055 !important;
    border-radius: 8px !important;
    color: #e8f0ff !important;
    font-size: 0.92rem !important;
}
.stTextInput > div > div > input:focus,
.stTextArea > div > div > textarea:focus {
    border-color: #f5c842 !important;
    box-shadow: 0 0 0 2px rgba(245,200,66,0.15) !important;
}

/* ── Expander ── */
.streamlit-expanderHeader {
    background: #0f1a2e !important;
    border: 1px solid #1e3055 !important;
    border-radius: 8px !important;
    color: #c8d8ec !important;
    font-weight: 600 !important;
}
.streamlit-expanderContent {
    background: #0d1631 !important;
    border: 1px solid #1e3055 !important;
    border-top: none !important;
}

/* ── Progress bar ── */
.stProgress > div > div > div > div {
    background: linear-gradient(90deg, #1e4494, #f5c842) !important;
    border-radius: 4px !important;
}

/* ── Success / info / warning ── */
.stSuccess {
    background: rgba(245,200,66,0.08) !important;
    border: 1px solid rgba(245,200,66,0.3) !important;
    border-radius: 8px !important;
    color: #f5c842 !important;
}
.stInfo {
    background: rgba(59,130,246,0.08) !important;
    border: 1px solid rgba(59,130,246,0.3) !important;
    border-radius: 8px !important;
}

/* ── Divider ── */
hr {
    border-color: #1e3055 !important;
    margin: 20px 0 !important;
}

/* ── Dataframe tables ── */
[data-testid="stDataFrame"] {
    border: 1px solid #1e3055 !important;
    border-radius: 10px !important;
    overflow: hidden;
}

/* ── Section labels ── */
.stSubheader, h2, h3 {
    color: #e8f0ff !important;
}

/* ── Sidebar logo area ── */
.sidebar-logo {
    text-align: center;
    padding: 16px 0 8px 0;
}
.sidebar-logo-text {
    font-size: 1.3rem;
    font-weight: 800;
    background: linear-gradient(90deg, #f5c842, #60a5fa);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    background-clip: text;
    letter-spacing: -0.5px;
}
.sidebar-logo-sub {
    font-size: 0.72rem;
    color: #4a6080;
    letter-spacing: 2px;
    text-transform: uppercase;
}
</style>
""", unsafe_allow_html=True)


# ─────────────────────────────────────────────────────────────────────────────
# Document History Page
# ─────────────────────────────────────────────────────────────────────────────

def _render_history_page():
    """Full document history page — list view with expand-on-click detail."""

    # ── CSS for history cards ─────────────────────────────────────────────────
    # History page uses global CSS — no local override needed
    pass

    # ── Fetch history ─────────────────────────────────────────────────────────
    with st.spinner("Loading history..."):
        res = requests.get(f"{BACKEND}/documents")

    if res.status_code != 200:
        st.error(f"Could not load history: {res.text}")
        return

    docs = res.json().get("documents", [])

    if not docs:
        st.info("No documents saved yet. Generate a document to see it here.")
        return

    st.markdown(f'<p style="color:#7a94b8;font-size:0.88rem;margin-bottom:8px;">{len(docs)} document(s) in your history</p>', unsafe_allow_html=True)
    st.markdown("---")

    for doc in docs:
        doc_id     = doc["id"]
        title      = doc["title"]
        doc_format = doc.get("doc_format", "word")
        created_at = doc.get("created_at", "")
        file_ext   = doc.get("file_ext", "docx")
        badge_cls  = "format-badge-excel" if doc_format == "excel" else "format-badge-word"
        badge_lbl  = "📊 Excel" if doc_format == "excel" else "📄 Word"

        # ── Card header ───────────────────────────────────────────────────────
        st.markdown(f"""
        <div class="hist-card">
            <span class="hist-title">{title}</span>
            &nbsp;&nbsp;<span class="{badge_cls}">{badge_lbl}</span>
            <div class="hist-meta">🕒 {created_at} &nbsp;|&nbsp; ID #{doc_id}</div>
        </div>
        """, unsafe_allow_html=True)

        # ── Action row ────────────────────────────────────────────────────────
        col1, col2, col3 = st.columns([2, 2, 8])

        with col1:
            view_key = f"view_{doc_id}"
            if st.button("👁 View", key=f"btn_view_{doc_id}", use_container_width=True):
                st.session_state[view_key] = not st.session_state.get(view_key, False)

        with col2:
            if st.button("🗑 Delete", key=f"btn_del_{doc_id}", use_container_width=True):
                del_res = requests.delete(f"{BACKEND}/documents/{doc_id}")
                if del_res.status_code == 200:
                    st.success(f"Deleted '{title}'")
                    st.rerun()
                else:
                    st.error("Delete failed")

        # ── Expanded detail (only when View is toggled on) ────────────────────
        if st.session_state.get(f"view_{doc_id}", False):
            with st.spinner("Loading document..."):
                detail_res = requests.get(f"{BACKEND}/documents/{doc_id}")

            if detail_res.status_code != 200:
                st.error("Could not load document.")
            else:
                detail = detail_res.json()
                doc_content = detail.get("content", {})

                with st.expander(f"📄 {title} — Full Content", expanded=True):

                    if doc_format == "excel":
                        sheets = doc_content.get("sheets", [])
                        for sheet in sheets:
                            sname   = sheet.get("sheet_name", "Sheet")
                            headers = sheet.get("headers", [])
                            rows    = sheet.get("rows", [])
                            st.markdown(f"**{sname}**")
                            if headers and rows:
                                padded = [r + [""] * max(0, len(headers) - len(r)) for r in rows]
                                st.dataframe(
                                    pd.DataFrame(padded, columns=headers),
                                    use_container_width=True,
                                    hide_index=True
                                )
                    else:
                        # Word document — render section by section
                        for section_name, section_content in doc_content.items():
                            st.markdown(f"**{section_name}**")
                            if isinstance(section_content, str):
                                st.write(section_content)
                            elif isinstance(section_content, list):
                                for item in section_content:
                                    st.write(f"• {item}")
                            elif isinstance(section_content, dict):
                                for k, v in section_content.items():
                                    st.write(f"**{k}:** {v}")
                            st.markdown("---")

                    # ── Download button ───────────────────────────────────────
                    dl_res = requests.get(
                        f"{BACKEND}/documents/{doc_id}/download",
                        stream=True
                    )
                    safe_title = title.replace(" ", "_")[:40]
                    hcol1, hcol2 = st.columns([1, 1])

                    with hcol1:
                        if dl_res.status_code == 200:
                            mime = (
                                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                                if file_ext == "xlsx"
                                else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                            st.download_button(
                                label=f"⬇️ Download .{file_ext}",
                                data=dl_res.content,
                                file_name=f"{safe_title}.{file_ext}",
                                mime=mime,
                                use_container_width=True,
                                key=f"dl_{doc_id}"
                            )
                        else:
                            st.caption("No file stored for download.")

                    with hcol2:
                        if st.button("🔗 Push to Notion", key=f"notion_{doc_id}",
                                     use_container_width=True):
                            push_to_notion_ui(
                                title, doc_format, doc_content,
                                db_id=doc_id,
                                notion_key=f"notion_hist_{doc_id}"
                            )

        st.markdown("")  # spacer between cards


def push_to_notion_ui(title: str, doc_format: str, doc_content: dict,
                      db_id: int = None, notion_key: str = "notion_page_id"):
    """Call /notion/push and store the returned page_id in session state."""
    with st.spinner("Pushing to Notion..."):
        res = requests.post(f"{BACKEND}/notion/push", json={
            "title":      title,
            "doc_format": doc_format,
            "content":    doc_content,
            "db_id":      db_id,
        })
    if res.status_code == 200:
        data = res.json()
        st.session_state[notion_key] = data.get("page_id")
        notion_url = data.get("url", "")
        st.success("✅ Pushed to Notion!")
        st.markdown(
            f'<a href="{notion_url}" target="_blank" style="'
            f'display:inline-block;margin-top:8px;padding:8px 18px;'
            f'background:linear-gradient(135deg,#f5c842,#e6b830);'
            f'color:#0a0e1a;font-weight:700;border-radius:8px;'
            f'text-decoration:none;font-size:0.88rem;">'
            f'🔗 Open in Notion</a>',
            unsafe_allow_html=True
        )
    else:
        st.error(f"Notion push failed: {res.text[:200]}")

# ── Sidebar navigation ───────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("""
    <div class="sidebar-logo">
        <div class="sidebar-logo-text">⚡ DocForge AI</div>
        <div class="sidebar-logo-sub">Document Generator</div>
    </div>
    """, unsafe_allow_html=True)
    st.markdown("---")
    page = st.radio(
        "Navigate",
        ["✏️  Generate Document", "📚  Document History"],
        label_visibility="collapsed"
    )
    st.markdown("---")
    st.markdown("""
    <div style="padding:12px;background:rgba(245,200,66,0.06);border:1px solid rgba(245,200,66,0.2);
                border-radius:8px;font-size:0.78rem;color:#7a94b8;line-height:1.6;">
        💾 Documents are <strong style="color:#f5c842">auto-saved</strong> to your database after generation.
    </div>
    """, unsafe_allow_html=True)

# ── Page routing ──────────────────────────────────────────────────────────────
if page == "📚  Document History":
    st.markdown("""
    <div class="page-header">
        <h1>📚 Document History</h1>
        <p>All your generated documents — browse, preview and download.</p>
    </div>
    """, unsafe_allow_html=True)
    _render_history_page()
    st.stop()

st.markdown("""
<div class="page-header">
    <h1>⚡ DocForge AI</h1>
    <p>Generate professional Word documents and Excel spreadsheets using AI — in seconds.</p>
</div>
""", unsafe_allow_html=True)


# ─────────────────────────────────────────────────────────────────────────────
# Helpers — Word rendering
# ─────────────────────────────────────────────────────────────────────────────

def render_section_content(content):
    if content is None or content == "":
        st.markdown('<p class="doc-paragraph"><em>No content provided.</em></p>', unsafe_allow_html=True)
        return
    if isinstance(content, str):
        content = content.strip()
        if content.startswith("{") or content.startswith("["):
            import json
            try:
                parsed = json.loads(content)
                render_section_content(parsed)
                return
            except Exception:
                pass
        for para in [p.strip() for p in content.split("\n") if p.strip()]:
            if para.startswith(("-", "*", "•", "·")):
                st.markdown(f'<p class="bullet-item">{para.lstrip("-*•· ").strip()}</p>', unsafe_allow_html=True)
            else:
                st.markdown(f'<p class="doc-paragraph">{para}</p>', unsafe_allow_html=True)
    elif isinstance(content, list):
        for item in content:
            if isinstance(item, dict):
                render_section_content(item)
            else:
                st.markdown(f'<p class="bullet-item">{str(item).strip().lstrip("-*•· ")}</p>', unsafe_allow_html=True)
    elif isinstance(content, dict):
        for key, val in content.items():
            st.markdown(f'<p class="doc-section-heading" style="font-size:1rem;margin-top:14px">{str(key).replace("_"," ").title()}</p>', unsafe_allow_html=True)
            render_section_content(val)
    else:
        st.markdown(f'<p class="doc-paragraph">{str(content)}</p>', unsafe_allow_html=True)


def render_full_document(title, sections):
    st.markdown('<div class="doc-card">', unsafe_allow_html=True)
    st.markdown(f'<p class="doc-title">📄 {title}</p>', unsafe_allow_html=True)
    for section_name, content in sections.items():
        st.markdown(f'<p class="doc-section-heading">{str(section_name).replace("_"," ").title()}</p>', unsafe_allow_html=True)
        render_section_content(content)
    st.markdown('</div>', unsafe_allow_html=True)


def flatten_to_text(content) -> str:
    if content is None: return ""
    if isinstance(content, str): return content.strip()
    if isinstance(content, list): return "\n".join(f"• {flatten_to_text(i)}" for i in content)
    if isinstance(content, dict):
        return "\n\n".join(f"{str(k).replace('_',' ').title()}:\n{flatten_to_text(v)}" for k, v in content.items())
    return str(content)


def create_word_document(title, sections_dict) -> BytesIO:
    doc = Document()
    doc.add_heading(title, 0).alignment = 1
    for section_name, content in sections_dict.items():
        doc.add_heading(str(section_name).replace("_", " ").title(), level=1)
        for para in [p.strip() for p in flatten_to_text(content).split("\n") if p.strip()]:
            if para.startswith("•"):
                doc.add_paragraph(para.lstrip("• ").strip(), style="List Bullet")
            else:
                p = doc.add_paragraph()
                p.add_run(para).font.size = Pt(11)
        doc.add_paragraph()
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─────────────────────────────────────────────────────────────────────────────
# Helpers — Excel rendering
# ─────────────────────────────────────────────────────────────────────────────

def render_excel_document(title, sheets):
    """Render each sheet as a styled DataFrame table."""
    st.markdown(f'<div class="doc-card"><p class="doc-title">📊 {title}</p></div>', unsafe_allow_html=True)

    for sheet in sheets:
        sheet_name  = sheet.get("sheet_name", "Sheet")
        description = sheet.get("description", "")
        headers     = sheet.get("headers", [])
        rows        = sheet.get("rows", [])
        header_rows = set(sheet.get("header_rows", []))
        totals_rows = set(sheet.get("totals_rows", []))
        notes       = sheet.get("notes", "")

        st.markdown(f'<p class="doc-section-heading">📋 {sheet_name}</p>', unsafe_allow_html=True)
        if description:
            st.caption(description)

        if headers and rows:
            # Pad rows to header length
            padded = [r + [""] * max(0, len(headers) - len(r)) for r in rows]
            df = pd.DataFrame(padded, columns=headers)

            # Normalise bold indices — LLM sometimes returns strings ("0") not ints
            bold_indices = set()
            for x in list(header_rows or []) + list(totals_rows or []):
                try:
                    bold_indices.add(int(x))
                except (ValueError, TypeError):
                    pass

            def style_row(row):
                base = "background-color:#FFFFFF; color:#000000; font-size:13px;"
                if int(row.name) in bold_indices:
                    return [base + "font-weight:bold;" for _ in row]
                return [base + "font-weight:normal;" for _ in row]

            styled = (
                df.style
                .apply(style_row, axis=1)
                .set_table_styles([{
                    "selector": "thead th",
                    "props": [
                        ("background-color", "#FFFFFF"),
                        ("color", "#000000"),
                        ("font-weight", "bold"),
                        ("font-size", "13px"),
                        ("border-bottom", "2px solid #cccccc")
                    ]
                }])
            )
            st.dataframe(styled, use_container_width=True, hide_index=True)

        if notes:
            st.caption(f"📝 Notes: {notes}")

        st.markdown("---")


def create_excel_file_from_data(excel_data: dict) -> BytesIO:
    """Call FastAPI /export/excel endpoint — keeps all openpyxl logic server-side."""
    res = requests.post(f"{BACKEND}/export/excel", json=excel_data)
    if res.status_code == 200:
        # Validate we got actual xlsx bytes, not a JSON error
        content_type = res.headers.get("content-type", "")
        if "spreadsheet" in content_type or "octet-stream" in content_type:
            return BytesIO(res.content)
        # Got JSON back instead of file — surface the error
        raise RuntimeError(f"Server returned unexpected response: {res.text[:300]}")
    raise RuntimeError(f"Excel export failed [{res.status_code}]: {res.text[:300]}")

def save_to_db(title: str, doc_type: str, doc_format: str,
               content: dict, file_buf: BytesIO = None, file_ext: str = None):
    """Save document to PostgreSQL via FastAPI. Shows success/error in UI."""
    payload = {
        "title":      title,
        "doc_type":   doc_type,
        "doc_format": doc_format,
        "content":    content,
        "file_ext":   file_ext,
    }
    if file_buf:
        file_buf.seek(0)
        payload["file_bytes"] = file_buf.read().hex()

    res = requests.post(f"{BACKEND}/documents/save", json=payload)
    if res.status_code == 200:
        doc_id = res.json().get("id")
        st.toast(f"✅ Saved to history (ID #{doc_id})", icon="💾")
        return doc_id
    else:
        st.warning(f"Could not save to history: {res.text[:200]}")
        return None




# ─────────────────────────────────────────────────────────────────────────────
# STEP 1: Plan
# ─────────────────────────────────────────────────────────────────────────────

st.markdown('<div class="step-label">Step 1 — Describe Your Document</div>', unsafe_allow_html=True)
user_prompt = st.text_input(
    "Describe the document you want to generate",
    placeholder="e.g. Balance Sheet for FY2024  •  SOP for lead qualification  •  Product Proposal",
    label_visibility="collapsed"
)

if st.button("⚡ Plan Document"):
    if not user_prompt.strip():
        st.warning("Please enter a document description.")
    else:
        with st.spinner("Planning document structure..."):
            res = requests.post(f"{BACKEND}/plan", json={"prompt": user_prompt})
        if res.status_code == 200:
            plan = res.json()
            st.session_state["title"]      = plan["title"]
            st.session_state["sections"]   = plan["sections"]
            st.session_state["doc_format"] = plan.get("doc_format", "word")
            for k in ["questions", "generated_sections", "excel_data"]:
                st.session_state.pop(k, None)
            st.success("Document structure generated!")
        else:
            st.error(f"Error: {res.text}")


# ─────────────────────────────────────────────────────────────────────────────
# STEP 2: Show Plan
# ─────────────────────────────────────────────────────────────────────────────

if "sections" in st.session_state:
    doc_format = st.session_state.get("doc_format", "word")
    badge = "excel" if doc_format == "excel" else "word"
    label = "📊 Excel / Spreadsheet" if doc_format == "excel" else "📄 Word Document"

    st.markdown('<div class="step-label">Step 2 — Review Structure</div>', unsafe_allow_html=True)
    col1, col2 = st.columns([3, 1])
    with col1:
        st.markdown(f"**Title:** {st.session_state['title']}")
    with col2:
        st.markdown(f'<span class="format-badge-{badge}">{label}</span>', unsafe_allow_html=True)

    for i, sec in enumerate(st.session_state["sections"]):
        col_sec, col_del = st.columns([10, 1])
        with col_sec:
            st.write(f"- {sec}")
        with col_del:
            if st.button("✕", key=f"del_sec_{i}", help=f"Remove '{sec}'"):
                st.session_state["sections"].pop(i)
                # Reset downstream so questions regenerate with new sections
                for k in ["questions", "generated_sections", "excel_data"]:
                    st.session_state.pop(k, None)
                st.rerun()

    # ── Add new section inline ────────────────────────────────────────────────
    with st.expander("➕ Add a section", expanded=False):
        new_sec = st.text_input(
            "Section name",
            placeholder="e.g. Risk Analysis, Appendix, Executive Summary...",
            key="new_section_input"
        )
        insert_pos = st.selectbox(
            "Insert position",
            options=["At the end"] + [f"Before: {s}" for s in st.session_state["sections"]],
            key="new_section_pos"
        )
        if st.button("Add Section", key="add_section_btn"):
            if new_sec.strip():
                sec_name = new_sec.strip()
                if insert_pos == "At the end":
                    st.session_state["sections"].append(sec_name)
                else:
                    ref = insert_pos.replace("Before: ", "")
                    idx = st.session_state["sections"].index(ref)
                    st.session_state["sections"].insert(idx, sec_name)
                # Reset downstream
                for k in ["questions", "generated_sections", "excel_data"]:
                    st.session_state.pop(k, None)
                st.success(f"Added '{sec_name}'. Regenerate questions to include it.")
                st.rerun()
            else:
                st.warning("Please enter a section name.")

    if st.button("📋 Generate Questions"):
        with st.spinner("Generating questions..."):
            res = requests.post(f"{BACKEND}/questions", json={
                "title": st.session_state["title"],
                "sections": st.session_state["sections"]
            })
        if res.status_code == 200:
            st.session_state["questions"] = res.json()
            for k in ["generated_sections", "excel_data"]:
                st.session_state.pop(k, None)
            st.success("Questions generated!")
        else:
            st.error(f"Error: {res.text}")


# ─────────────────────────────────────────────────────────────────────────────
# STEP 3: Collect Answers
# ─────────────────────────────────────────────────────────────────────────────

if "questions" in st.session_state:
    st.markdown('<div class="step-label">Step 3 — Provide Details (optional)</div>', unsafe_allow_html=True)
    answers = {}
    for section, qs in st.session_state["questions"].items():
        # Skip any non-list values that may have been injected (e.g. _cached: True)
        if not isinstance(qs, list):
            continue
        st.markdown(f"### {section}")
        for q in qs:
            answers[q] = st.text_area(q, key=f"{section}__{q}")

    # ── Add section at Step 3 ─────────────────────────────────────────────────
    with st.expander("➕ Add another section", expanded=False):
        new_sec3 = st.text_input(
            "New section name",
            placeholder="e.g. Competitive Analysis, Budget Breakdown...",
            key="new_section_input_s3"
        )
        insert_pos3 = st.selectbox(
            "Insert position",
            options=["At the end"] + [f"Before: {s}" for s in st.session_state.get("sections", [])],
            key="new_section_pos_s3"
        )
        if st.button("Add & Regenerate Questions", key="add_section_btn_s3"):
            if new_sec3.strip():
                sec_name3 = new_sec3.strip()
                sections = st.session_state.get("sections", [])
                if insert_pos3 == "At the end":
                    sections.append(sec_name3)
                else:
                    ref3 = insert_pos3.replace("Before: ", "")
                    idx3 = sections.index(ref3)
                    sections.insert(idx3, sec_name3)
                st.session_state["sections"] = sections

                # Regenerate questions for all sections including the new one
                with st.spinner(f"Generating questions for '{sec_name3}'..."):
                    res = requests.post(f"{BACKEND}/questions", json={
                        "title":    st.session_state["title"],
                        "sections": sections
                    })
                if res.status_code == 200:
                    st.session_state["questions"] = res.json()
                    st.success(f"Added '{sec_name3}' and regenerated questions.")
                    st.rerun()
                else:
                    st.error(f"Could not regenerate questions: {res.text}")
            else:
                st.warning("Please enter a section name.")

    if st.button("🚀 Generate Document"):
        with st.spinner("Starting generation job..."):
            res = requests.post(f"{BACKEND}/generate", json={
                "title":      st.session_state["title"],
                "sections":   st.session_state["sections"],
                "answers":    answers,
                "doc_format": st.session_state.get("doc_format", "word")
            })

        if res.status_code != 200:
            st.error(f"Error: {res.text}")
        else:
            raw = res.json()
            doc_format = st.session_state.get("doc_format", "word")
            if doc_format == "excel":
                # Normalise: strip top-level keys that aren't part of the sheet structure
                # LLM sometimes returns {"doc_type":"excel","title":"...","sheets":[...]}
                if isinstance(raw, dict) and "sheets" not in raw:
                    # Flat structure — find the sheets inside any nested key
                    for v in raw.values():
                        if isinstance(v, dict) and "sheets" in v:
                            raw = v
                            break
                # Ensure title is present
                if isinstance(raw, dict) and "title" not in raw:
                    raw["title"] = st.session_state.get("title", "Document")
                # Remove doc_type key — not needed downstream
                raw.pop("doc_type", None)
                st.session_state["excel_data"] = raw
                st.session_state.pop("generated_sections", None)
                # Auto-save excel doc to DB
                save_to_db(
                    title      = st.session_state["title"],
                    doc_type   = st.session_state["title"],
                    doc_format = "excel",
                    content    = st.session_state["excel_data"],
                    file_ext   = "xlsx"
                )
            else:
                # Unwrap any wrapper keys
                if isinstance(raw, dict) and list(raw.keys()) == ["Document"]:
                    raw = raw["Document"]
                if isinstance(raw, dict) and list(raw.keys()) == ["sections"]:
                    raw = raw["sections"]
                st.session_state["generated_sections"] = raw if isinstance(raw, dict) else {"Content": str(raw)}
                st.session_state.pop("excel_data", None)

            st.success("Document generated!")
            # Auto-save word doc (excel is saved in its own branch above)
            if st.session_state.get("doc_format", "word") == "word":
                save_to_db(
                    title      = st.session_state["title"],
                    doc_type   = st.session_state["title"],
                    doc_format = "word",
                    content    = st.session_state.get("generated_sections", {}),
                    file_ext   = "docx"
                )
            st.rerun()


# ─────────────────────────────────────────────────────────────────────────────
# STEP 4A: Display Word Document
# ─────────────────────────────────────────────────────────────────────────────

if "generated_sections" in st.session_state:
    st.divider()
    st.markdown('<div class="step-label">Generated Document</div>', unsafe_allow_html=True)
    render_full_document(st.session_state["title"], st.session_state["generated_sections"])

    st.divider()
    st.markdown('<div class="step-label">Review & Refine</div>', unsafe_allow_html=True)

    for section, content in list(st.session_state["generated_sections"].items()):
        with st.expander(f"✏️ Edit: {section}", expanded=False):
            feedback = st.text_area(f"Suggest changes for '{section}'", key=f"feedback_{section}")
            if st.button(f"Update '{section}'", key=f"btn_{section}"):
                with st.spinner(f"Refining {section}..."):
                    res = requests.post(f"{BACKEND}/refine-section", json={
                        "section_name":  section,
                        "original_text": flatten_to_text(content),
                        "feedback":      feedback,
                        "doc_format":    "word"
                    })
                if res.status_code == 200:
                    result = res.json()
                    updated = result.get("updated_text", "")
                    if isinstance(updated, (dict, list)):
                        updated = flatten_to_text(updated)
                    st.session_state["generated_sections"] = {
                        **st.session_state["generated_sections"],
                        section: str(updated)
                    }
                    st.success(f"'{section}' updated!")
                    st.rerun()
                elif res.status_code == 429:
                    st.warning("⏳ Too many refinements — please wait a moment before trying again.")
                else:
                    st.error(f"Error: {res.text}")

    st.divider()
    st.markdown('<div class="step-label">Download & Export</div>', unsafe_allow_html=True)
    docx_file  = create_word_document(st.session_state["title"], st.session_state["generated_sections"])
    safe_title = st.session_state["title"].replace(" ", "_")[:40]

    col_dl, col_notion = st.columns([1, 1])
    with col_dl:
        st.download_button(
            label="⬇️ Download Word (.docx)",
            data=docx_file,
            file_name=f"{safe_title}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
    with col_notion:
        already_pushed = st.session_state.get("notion_page_id_word")
        btn_label = "🔄 Update in Notion" if already_pushed else "🔗 Push to Notion"
        if st.button(btn_label, key="notion_word_btn", use_container_width=True):
            if already_pushed:
                # Update existing page
                with st.spinner("Updating Notion page..."):
                    res = requests.post(f"{BACKEND}/notion/update", json={
                        "page_id":    already_pushed,
                        "title":      st.session_state["title"],
                        "doc_format": "word",
                        "content":    st.session_state["generated_sections"],
                    })
                if res.status_code == 200:
                    notion_url = res.json().get("url", "")
                    st.success("✅ Notion page updated!")
                    st.markdown(f'<a href="{notion_url}" target="_blank" style="display:inline-block;margin-top:8px;padding:8px 18px;background:linear-gradient(135deg,#f5c842,#e6b830);color:#0a0e1a;font-weight:700;border-radius:8px;text-decoration:none;font-size:0.88rem;">🔗 Open in Notion</a>', unsafe_allow_html=True)
                else:
                    st.error(f"Update failed: {res.text[:200]}")
            else:
                push_to_notion_ui(
                    st.session_state["title"], "word",
                    st.session_state["generated_sections"],
                    notion_key="notion_page_id_word"
                )


# ─────────────────────────────────────────────────────────────────────────────
# STEP 4B: Display Excel Document
# ─────────────────────────────────────────────────────────────────────────────

if "excel_data" in st.session_state:
    excel_data = st.session_state["excel_data"]
    sheets     = excel_data.get("sheets", [])

    st.divider()
    st.markdown('<div class="step-label">Generated Spreadsheet</div>', unsafe_allow_html=True)
    render_excel_document(st.session_state["title"], sheets)

    # ── Per-sheet refinement ──────────────────────────────────────────────────
    st.divider()
    st.markdown('<div class="step-label">Review & Refine Sheets</div>', unsafe_allow_html=True)

    for i, sheet in enumerate(sheets):
        sheet_name = sheet.get("sheet_name", f"Sheet {i+1}")
        with st.expander(f"✏️ Edit: {sheet_name}", expanded=False):

            # Show current data as readable table preview
            headers = sheet.get("headers", [])
            rows    = sheet.get("rows", [])
            if headers and rows:
                padded = [r + [""] * max(0, len(headers) - len(r)) for r in rows]
                st.dataframe(
                    pd.DataFrame(padded, columns=headers),
                    use_container_width=True,
                    hide_index=True
                )

            feedback = st.text_area(
                f"Suggest changes for '{sheet_name}'",
                placeholder="Be specific. e.g. Update Cash value to 250000 • Add row: Deferred Tax = 75000 • Change description to none",
                key=f"excel_feedback_{i}",
                height=100
            )

            if st.button(f"Update '{sheet_name}'", key=f"excel_btn_{i}"):
                with st.spinner(f"Refining {sheet_name}..."):
                    res = requests.post(f"{BACKEND}/refine-section", json={
                        "section_name": sheet_name,
                        "current_data": sheet,      # pass full sheet so LLM keeps all rows
                        "feedback":     feedback,
                        "doc_format":   "excel"
                    })

                if res.status_code == 200:
                    result = res.json()
                    updated_sheet = result.get("updated_sheet", sheet)

                    # Validate the updated sheet has rows before accepting it
                    if not updated_sheet.get("rows"):
                        st.error("Refinement returned empty rows — keeping original. Try rephrasing your feedback.")
                    else:
                        # Surgically replace only this sheet, preserve all others
                        updated_sheets = list(sheets)   # copy current list
                        updated_sheets[i] = updated_sheet
                        st.session_state["excel_data"] = {
                            **st.session_state["excel_data"],
                            "sheets": updated_sheets
                        }
                        st.success(f"'{sheet_name}' updated!")
                        st.rerun()
                else:
                    st.error(f"Refinement failed [{res.status_code}]: {res.text[:300]}")

    # ── Download Excel ────────────────────────────────────────────────────────
    st.divider()
    st.markdown('<div class="step-label">Download & Export</div>', unsafe_allow_html=True)

    try:
        xlsx_file  = create_excel_file_from_data(st.session_state["excel_data"])
        safe_title = st.session_state["title"].replace(" ", "_")[:40]

        col_dl, col_notion = st.columns([1, 1])
        with col_dl:
            st.download_button(
                label="⬇️ Download Excel (.xlsx)",
                data=xlsx_file,
                file_name=f"{safe_title}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True
            )
        with col_notion:
            already_pushed = st.session_state.get("notion_page_id_excel")
            btn_label = "🔄 Update in Notion" if already_pushed else "🔗 Push to Notion"
            if st.button(btn_label, key="notion_excel_btn", use_container_width=True):
                if already_pushed:
                    with st.spinner("Updating Notion page..."):
                        res = requests.post(f"{BACKEND}/notion/update", json={
                            "page_id":    already_pushed,
                            "title":      st.session_state["title"],
                            "doc_format": "excel",
                            "content":    st.session_state["excel_data"],
                        })
                    if res.status_code == 200:
                        notion_url = res.json().get("url", "")
                        st.success("✅ Notion page updated!")
                        st.markdown(f'<a href="{notion_url}" target="_blank" style="display:inline-block;margin-top:8px;padding:8px 18px;background:linear-gradient(135deg,#f5c842,#e6b830);color:#0a0e1a;font-weight:700;border-radius:8px;text-decoration:none;font-size:0.88rem;">🔗 Open in Notion</a>', unsafe_allow_html=True)
                    else:
                        st.error(f"Update failed: {res.text[:200]}")
                else:
                    push_to_notion_ui(
                        st.session_state["title"], "excel",
                        st.session_state["excel_data"],
                        notion_key="notion_page_id_excel"
                    )
    except Exception as e:
        st.error(f"Excel export error: {e}")
        st.info("Make sure `openpyxl` is installed: `pip install openpyxl`")