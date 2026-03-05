mport streamlit as st
import requests
from docx import Document
from docx.shared import Pt, RGBColor
from io import BytesIO

BACKEND = "http://127.0.0.1:8000"

st.set_page_config(page_title="Agentic Document Generator", layout="wide")

# ─────────────────────────────────────────────────────────────────────────────
# Custom CSS
# ─────────────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
.doc-title {
    font-size: 2rem;
    font-weight: 700;
    color: #e8eaf0;
    border-bottom: 2px solid #4f8ef7;
    padding-bottom: 10px;
    margin-bottom: 24px;
}
.doc-section-heading {
    font-size: 1.2rem;
    font-weight: 600;
    color: #4f8ef7;
    margin-top: 28px;
    margin-bottom: 8px;
    padding-left: 10px;
    border-left: 3px solid #4f8ef7;
}
.doc-paragraph {
    font-size: 0.97rem;
    color: #d0d3e0;
    line-height: 1.8;
    margin-bottom: 10px;
    padding-left: 13px;
}
.doc-card {
    background: #1a1d27;
    border: 1px solid #2d3142;
    border-radius: 12px;
    padding: 28px 32px;
    margin-bottom: 20px;
}
.bullet-item {
    color: #d0d3e0;
    font-size: 0.95rem;
    line-height: 1.8;
    padding-left: 24px;
    position: relative;
}
.bullet-item::before {
    content: "•";
    color: #4f8ef7;
    position: absolute;
    left: 8px;
}
</style>
""", unsafe_allow_html=True)

st.title("🧠 Agentic Document Generator")


# ─────────────────────────────────────────────────────────────────────────────
# Helper: Render section content properly
# ─────────────────────────────────────────────────────────────────────────────

def render_section_content(content):
    """
    Intelligently renders section content as formatted text.
    Handles: plain strings, dicts, lists, nested structures, JSON blobs.
    """
    if content is None or content == "":
        st.markdown('<p class="doc-paragraph"><em>No content provided.</em></p>', unsafe_allow_html=True)
        return

    # ── If it's a plain string ────────────────────────────────────────────────
    if isinstance(content, str):
        content = content.strip()

        # Try to parse if it looks like JSON
        if content.startswith("{") or content.startswith("["):
            import json
            try:
                parsed = json.loads(content)
                render_section_content(parsed)  # recurse with parsed value
                return
            except Exception:
                pass  # not JSON, treat as plain text

        # Render as paragraphs (split on double newlines or single newlines)
        paragraphs = [p.strip() for p in content.split("\n") if p.strip()]
        for para in paragraphs:
            # Detect bullet-like lines
            if para.startswith(("-", "*", "•", "·")):
                clean = para.lstrip("-*•· ").strip()
                st.markdown(f'<p class="bullet-item">{clean}</p>', unsafe_allow_html=True)
            elif para.startswith(tuple("0123456789")) and (para[1] in ".)" or para[2:3] in ".):"):
                # Numbered list item
                st.markdown(f'<p class="bullet-item">{para}</p>', unsafe_allow_html=True)
            else:
                st.markdown(f'<p class="doc-paragraph">{para}</p>', unsafe_allow_html=True)

    # ── If it's a list ────────────────────────────────────────────────────────
    elif isinstance(content, list):
        for item in content:
            if isinstance(item, dict):
                render_section_content(item)
            elif isinstance(item, str):
                clean = item.strip().lstrip("-*•· ")
                st.markdown(f'<p class="bullet-item">{clean}</p>', unsafe_allow_html=True)
            else:
                st.markdown(f'<p class="bullet-item">{str(item)}</p>', unsafe_allow_html=True)

    # ── If it's a dict ────────────────────────────────────────────────────────
    elif isinstance(content, dict):
        for key, val in content.items():
            label = str(key).replace("_", " ").title()
            st.markdown(f'<p class="doc-section-heading" style="font-size:1rem;margin-top:14px">{label}</p>', unsafe_allow_html=True)
            if isinstance(val, (list, dict)):
                render_section_content(val)
            else:
                clean = str(val).strip()
                if clean:
                    st.markdown(f'<p class="doc-paragraph">{clean}</p>', unsafe_allow_html=True)

    # ── Fallback ──────────────────────────────────────────────────────────────
    else:
        st.markdown(f'<p class="doc-paragraph">{str(content)}</p>', unsafe_allow_html=True)


def render_full_document(title: str, sections: dict):
    """Renders the complete document in a styled card."""
    st.markdown(f'<div class="doc-card">', unsafe_allow_html=True)
    st.markdown(f'<p class="doc-title">📄 {title}</p>', unsafe_allow_html=True)

    for section_name, content in sections.items():
        heading = str(section_name).replace("_", " ").title()
        st.markdown(f'<p class="doc-section-heading">{heading}</p>', unsafe_allow_html=True)
        render_section_content(content)

    st.markdown('</div>', unsafe_allow_html=True)


# ─────────────────────────────────────────────────────────────────────────────
# Helper: Create DOCX
# ─────────────────────────────────────────────────────────────────────────────

def flatten_to_text(content) -> str:
    """Recursively flatten any content type to a plain text string."""
    if content is None:
        return ""
    if isinstance(content, str):
        return content.strip()
    if isinstance(content, list):
        return "\n".join(f"• {flatten_to_text(item)}" for item in content)
    if isinstance(content, dict):
        parts = []
        for k, v in content.items():
            label = str(k).replace("_", " ").title()
            parts.append(f"{label}:\n{flatten_to_text(v)}")
        return "\n\n".join(parts)
    return str(content)


def create_word_document(title: str, sections_dict: dict) -> BytesIO:
    doc = Document()

    # Title
    title_para = doc.add_heading(title, 0)
    title_para.alignment = 1  # Center

    for section_name, content in sections_dict.items():
        heading = str(section_name).replace("_", " ").title()
        doc.add_heading(heading, level=1)

        text = flatten_to_text(content)
        paragraphs = [p.strip() for p in text.split("\n") if p.strip()]

        for para in paragraphs:
            if para.startswith("•"):
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(para.lstrip("• ").strip())
            else:
                p = doc.add_paragraph()
                run = p.add_run(para)
                run.font.size = Pt(11)

        doc.add_paragraph()  # spacer

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


# ─────────────────────────────────────────────────────────────────────────────
# STEP 1: User Prompt → Plan
# ─────────────────────────────────────────────────────────────────────────────

user_prompt = st.text_input("Describe the document you want to generate",
                            placeholder="e.g. SOP for enterprise lead qualification in a SaaS company")

if st.button("Plan Document"):
    if not user_prompt.strip():
        st.warning("Please enter a document description.")
    else:
        with st.spinner("Planning document structure..."):
            res = requests.post(f"{BACKEND}/plan", json={"prompt": user_prompt})
        if res.status_code == 200:
            plan = res.json()
            st.session_state["title"] = plan["title"]
            st.session_state["sections"] = plan["sections"]
            # Reset downstream state
            for key in ["questions", "generated_sections"]:
                st.session_state.pop(key, None)
            st.success("Document structure generated!")
        else:
            st.error(f"Error: {res.text}")


# ─────────────────────────────────────────────────────────────────────────────
# STEP 2: Show Plan → Generate Questions
# ─────────────────────────────────────────────────────────────────────────────

if "sections" in st.session_state:
    st.subheader("📋 Document Structure")
    st.markdown(f"**Title:** {st.session_state['title']}")
    for sec in st.session_state["sections"]:
        st.write(f"- {sec}")

    if st.button("Generate Questions"):
        with st.spinner("Generating questions for each section..."):
            res = requests.post(f"{BACKEND}/questions", json={
                "title": st.session_state["title"],
                "sections": st.session_state["sections"]
            })
        if res.status_code == 200:
            st.session_state["questions"] = res.json()
            st.session_state.pop("generated_sections", None)
            st.success("Questions generated!")
        else:
            st.error(f"Error: {res.text}")


# ─────────────────────────────────────────────────────────────────────────────
# STEP 3: Collect Answers → Generate Document
# ─────────────────────────────────────────────────────────────────────────────

if "questions" in st.session_state:
    st.subheader("📝 Provide Inputs")
    answers = {}

    for section, qs in st.session_state["questions"].items():
        st.markdown(f"### {section}")
        for q in qs:
            answers[q] = st.text_area(q, key=f"{section}__{q}")

    if st.button("Generate Document"):
        with st.spinner("Generating document content..."):
            res = requests.post(f"{BACKEND}/generate", json={
                "title": st.session_state["title"],
                "sections": st.session_state["sections"],
                "answers": answers
            })
        if res.status_code == 200:
            raw = res.json()

            # Normalize: backend might return {sections: {...}} or the dict directly
            if isinstance(raw, dict) and "sections" in raw:
                st.session_state["generated_sections"] = raw["sections"]
            elif isinstance(raw, dict):
                st.session_state["generated_sections"] = raw
            else:
                st.session_state["generated_sections"] = {"Content": str(raw)}

            st.success("Document generated successfully!")
            st.rerun()
        else:
            st.error(f"Error: {res.text}")


# ─────────────────────────────────────────────────────────────────────────────
# STEP 4: Display Document + Section Refinement
# ─────────────────────────────────────────────────────────────────────────────

if "generated_sections" in st.session_state:
    st.divider()
    st.subheader("📄 Generated Document")

    render_full_document(
        st.session_state["title"],
        st.session_state["generated_sections"]
    )

    # ── Section-wise Refinement ───────────────────────────────────────────────
    st.divider()
    st.subheader("🔍 Review & Edit Sections")

    for section, content in list(st.session_state["generated_sections"].items()):
        with st.expander(f"✏️ Edit: {section}", expanded=False):
            # Show current content as editable plain text
            current_text = flatten_to_text(content)
            feedback = st.text_area(
                f"Suggest changes for '{section}'",
                placeholder="e.g. Make it more formal, add ISO compliance reference...",
                key=f"feedback_{section}"
            )
            if st.button(f"Update '{section}'", key=f"btn_{section}"):
                with st.spinner(f"Refining {section}..."):
                    res = requests.post(f"{BACKEND}/refine-section", json={
                        "section_name": section,
                        "original_text": current_text,
                        "feedback": feedback
                    })
                if res.status_code == 200:
                    result = res.json()
                    # Handle various response shapes
                    if isinstance(result, dict) and "updated_text" in result:
                        updated = result["updated_text"]
                    elif isinstance(result, str):
                        updated = result
                    else:
                        updated = str(result)
                    st.session_state["generated_sections"][section] = updated
                    st.success(f"'{section}' updated!")
                    st.rerun()
                else:
                    st.error(f"Error: {res.text}")

    # ── Download DOCX ─────────────────────────────────────────────────────────
    st.divider()
    st.subheader("⬇️ Download Document")

    docx_file = create_word_document(
        st.session_state["title"],
        st.session_state["generated_sections"]
    )
    safe_title = st.session_state["title"].replace(" ", "_")[:40]
    st.download_button(
        label="⬇️ Download as Word Document (.docx)",
        data=docx_file,
        file_name=f"{safe_title}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )